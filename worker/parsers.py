"""
Document parsing utilities for extracting text from various file formats.

This module provides functions to parse different document types (TXT, PDF, DOCX)
and extract plain text, plus utilities for chunking text with overlap.
"""

from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from typing import List, Optional
import logging
import re

logger = logging.getLogger(__name__)

def parse_txt(path: Path) -> str:
    """
    Parse plain text file and return content.
    
    Args:
        path: Path to the text file
        
    Returns:
        str: Plain text content
        
    Raises:
        FileNotFoundError: If file doesn't exist
        UnicodeDecodeError: If file encoding issues occur
    """
    try:
        if not path.exists():
            raise FileNotFoundError(f"Text file not found: {path}")
        
        # Try UTF-8 first, fallback to other encodings
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decode failed for {path}, trying latin-1")
            content = path.read_text(encoding="latin-1", errors="ignore")
        
        logger.debug(f"Parsed text file: {path} ({len(content)} chars)")
        return content.strip()
        
    except Exception as e:
        logger.error(f"Error parsing text file {path}: {e}")
        raise

def parse_pdf(path: Path) -> str:
    """
    Parse PDF file and extract text from all pages.
    Handles corrupted PDFs and image-based PDFs with OCR fallback.
    
    Args:
        path: Path to the PDF file
        
    Returns:
        str: Extracted plain text from all pages
        
    Raises:
        FileNotFoundError: If file doesn't exist
        Exception: If PDF parsing fails completely
    """
    try:
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {path}")
        
        doc = None
        try:
            # Try to open the PDF
            doc = fitz.open(str(path))
            
            # Check if PDF is encrypted
            if doc.needs_pass:
                logger.warning(f"PDF {path} is password protected, skipping")
                if doc:
                    doc.close()
                return ""
            
            texts = []
            
            for page_num in range(doc.page_count):
                try:
                    page = doc[page_num]
                    page_text = page.get_text("text")
                    
                    # If no text found, try OCR on the page image
                    if not page_text.strip():
                        print(f"📄 PDF Page {page_num + 1}: No text found, attempting OCR...")
                        logger.info(f"No text found on page {page_num + 1} of {path}, trying OCR")
                        
                        try:
                            # Convert PDF page to image and run OCR
                            page_text = _extract_text_from_pdf_page_ocr(page, page_num + 1, path.name)
                            if page_text.strip():
                                print(f"✅ PDF OCR: Extracted {len(page_text)} characters from page {page_num + 1}")
                            else:
                                print(f"⚠️  PDF OCR: No text found on page {page_num + 1}")
                        except Exception as ocr_error:
                            print(f"❌ PDF OCR failed on page {page_num + 1}: {ocr_error}")
                            logger.warning(f"OCR failed for page {page_num + 1} of {path}: {ocr_error}")
                            continue
                    
                    # Clean up text formatting
                    if page_text.strip():
                        page_text = re.sub(r'\n+', '\n', page_text)  # Multiple newlines to single
                        page_text = re.sub(r' +', ' ', page_text)     # Multiple spaces to single
                        texts.append(page_text.strip())
                        
                except Exception as page_error:
                    logger.warning(f"Error processing page {page_num + 1} of {path}: {page_error}")
                    continue
            
            if doc:
                doc.close()
            
            if not texts:
                logger.warning(f"No extractable text found in PDF: {path}")
                return ""
            
            full_text = "\n\n".join(texts)
            logger.debug(f"Parsed PDF: {path} ({len(texts)} pages with text, {len(full_text)} chars)")
            return full_text
            
        except Exception as doc_error:
            if doc:
                doc.close()
            logger.error(f"Error opening PDF document {path}: {doc_error}")
            # Return empty string instead of raising exception
            return ""
        
    except Exception as e:
        logger.error(f"Error parsing PDF {path}: {e}")
        # Return empty string instead of raising exception to continue processing other files
        return ""

def parse_docx(path: Path) -> str:
    """
    Parse DOCX file and extract text from all paragraphs.
    
    Args:
        path: Path to the DOCX file
        
    Returns:
        str: Extracted plain text from all paragraphs
        
    Raises:
        FileNotFoundError: If file doesn't exist
        Exception: If DOCX parsing fails
    """
    try:
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        
        doc = Document(str(path))
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append(text)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n\n".join(paragraphs)
        logger.debug(f"Parsed DOCX: {path} ({len(doc.paragraphs)} paragraphs, {len(full_text)} chars)")
        return full_text
        
    except Exception as e:
        logger.error(f"Error parsing DOCX {path}: {e}")
        # Return empty string instead of raising exception
        return ""

def parse_image_ocr(path: Path) -> str:
    """
    Parse image file using OCR and extract text.
    
    Args:
        path: Path to the image file
        
    Returns:
        str: Extracted text from image, empty string if OCR fails
        
    Note:
        This function gracefully handles OCR failures and returns empty string
        rather than raising exceptions, as OCR can be unreliable.
    """
    try:
        if not path.exists():
            print(f"❌ OCR: Image file not found: {path}")
            logger.warning(f"Image file not found: {path}")
            return ""
        
        print(f"🖼️  OCR: Processing image {path.name}")
        logger.info(f"🔍 Starting OCR processing for image: {path}")
        
        from PIL import Image, ImageEnhance, ImageFilter
        import pytesseract
        
        # Get file size for reporting
        file_size = path.stat().st_size
        print(f"📏 OCR: Image size: {file_size} bytes")
        
        # Open and process image
        with Image.open(str(path)) as img:
            print(f"🎨 OCR: Image dimensions: {img.size[0]}x{img.size[1]} pixels, Mode: {img.mode}")
            
            # Convert to RGB if needed
            if img.mode != 'RGB':
                print(f"🔄 OCR: Converting image from {img.mode} to RGB")
                img = img.convert('RGB')
            
            # Try multiple OCR approaches for better handwriting recognition
            print(f"🤖 OCR: Extracting text using multiple configurations...")
            
            # Configuration 1: Standard OCR (current approach)
            config1 = r'--oem 3 --psm 6'
            text1 = pytesseract.image_to_string(img, lang='eng', config=config1)
            
            # Configuration 2: Better for handwriting - single text block
            config2 = r'--oem 1 --psm 7'  # LSTM engine, single text line
            text2 = pytesseract.image_to_string(img, lang='eng', config=config2)
            
            # Configuration 3: Enhanced preprocessing for handwriting
            enhanced_img = _enhance_image_for_ocr(img)
            config3 = r'--oem 3 --psm 3'  # Full page segmentation
            text3 = pytesseract.image_to_string(enhanced_img, lang='eng', config=config3)
            
            # Choose the best result based on length and content quality
            results = [
                ("Standard OCR", text1),
                ("Handwriting Mode", text2), 
                ("Enhanced Preprocessing", text3)
            ]
            
            # Select the longest non-empty result as it's likely the most complete
            best_text = ""
            best_method = "Standard OCR"
            for method, text in results:
                cleaned_text = text.strip()
                if len(cleaned_text) > len(best_text):
                    best_text = cleaned_text
                    best_method = method
            
            # Clean up OCR text
            text = re.sub(r'\n+', '\n', best_text)  # Multiple newlines to single
            text = re.sub(r' +', ' ', text)         # Multiple spaces to single
            text = text.strip()
            
            if text:
                print(f"✅ OCR: Successfully extracted {len(text)} characters using {best_method}")
                print(f"📝 OCR: Text preview: {text[:100]}{'...' if len(text) > 100 else ''}")
                logger.info(f"OCR success: {path.name} - extracted {len(text)} chars using {best_method}")
            else:
                print(f"⚠️  OCR: No text found in image with any method")
                logger.warning(f"OCR found no text in image: {path}")
            
            return text
            
    except ImportError as e:
        print(f"❌ OCR: Missing dependencies - PIL or pytesseract not available")
        logger.error(f"OCR dependencies missing: {e}")
        print("💡 OCR: Install with: pip install pillow pytesseract")
        return ""
    except Exception as e:
        print(f"❌ OCR: Processing failed for {path.name}: {str(e)}")
        logger.error(f"OCR failed for image {path}: {e}")
        return ""

def _enhance_image_for_ocr(img):
    """
    Enhance image for better OCR performance, especially for handwritten content.
    
    Args:
        img: PIL Image object
        
    Returns:
        Enhanced PIL Image object
    """
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        
        # Convert to grayscale for better OCR
        if img.mode != 'L':
            img = img.convert('L')
        
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        # Enhance sharpness
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(2.0)
        
        # Apply slight blur to reduce noise
        img = img.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
        
        # Scale up image for better recognition
        width, height = img.size
        new_size = (int(width * 1.5), int(height * 1.5))
        img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        return img
        
    except Exception as e:
        logger.warning(f"Image enhancement failed: {e}")
        return img

def _extract_text_from_pdf_page_ocr(page, page_num: int, pdf_name: str) -> str:
    """
    Extract text from a PDF page using OCR.
    
    Args:
        page: PyMuPDF page object
        page_num: Page number for logging
        pdf_name: PDF filename for logging
        
    Returns:
        str: Extracted text from the page
    """
    try:
        from PIL import Image
        import pytesseract
        import io
        
        # Convert PDF page to image
        # Use high resolution for better OCR
        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
        pix = page.get_pixmap(matrix=mat)
        
        # Convert to PIL Image
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        
        print(f"🔍 PDF OCR: Converting page {page_num} to image ({img.size[0]}x{img.size[1]} pixels)")
        
        # Apply the same enhancement as for regular images
        enhanced_img = _enhance_image_for_ocr(img)
        
        # Use multiple OCR configurations
        configs = [
            r'--oem 3 --psm 6',  # Standard
            r'--oem 1 --psm 3',  # LSTM with auto page segmentation  
            r'--oem 3 --psm 1'   # Auto page segmentation with OSD
        ]
        
        best_text = ""
        best_method = "Standard"
        
        for i, config in enumerate(configs):
            try:
                method_names = ["Standard", "LSTM Auto", "Auto with OSD"]
                text = pytesseract.image_to_string(enhanced_img, lang='eng', config=config)
                text = text.strip()
                
                if len(text) > len(best_text):
                    best_text = text
                    best_method = method_names[i]
                    
            except Exception as config_error:
                continue
        
        # Clean up the text
        if best_text:
            best_text = re.sub(r'\n+', '\n', best_text)
            best_text = re.sub(r' +', ' ', best_text)
            best_text = best_text.strip()
            
            print(f"🎯 PDF OCR: Best result from {best_method} method")
        
        return best_text
        
    except ImportError:
        logger.warning("PIL or pytesseract not available for PDF OCR")
        return ""
    except Exception as e:
        logger.warning(f"PDF page OCR failed: {e}")
        return ""

def chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    """
    Split text into chunks with word-level overlap.
    
    Args:
        text: Input text to chunk
        chunk_size: Maximum number of characters per chunk
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List[str]: List of text chunks, whitespace-trimmed and deterministic
        
    Raises:
        ValueError: If parameters are invalid
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap < 0:
        raise ValueError("overlap cannot be negative")
    if overlap >= chunk_size:
        raise ValueError("overlap must be less than chunk_size")
    
    # Clean and normalize text
    text = text.strip()
    if not text:
        return []
    
    # Split into words for better chunk boundaries
    words = text.split()
    if not words:
        return []
    
    chunks = []
    start_idx = 0
    
    while start_idx < len(words):
        # Build chunk by adding words until we exceed chunk_size
        chunk_words = []
        char_count = 0
        
        for i in range(start_idx, len(words)):
            word = words[i]
            # Account for space between words (except for first word)
            word_length = len(word) + (1 if chunk_words else 0)
            
            if char_count + word_length > chunk_size and chunk_words:
                # We've hit the limit, stop here
                break
            
            chunk_words.append(word)
            char_count += word_length
        
        # Create chunk text
        if chunk_words:
            chunk_text = " ".join(chunk_words).strip()
            chunks.append(chunk_text)
        
        # Calculate next starting position with overlap
        if overlap > 0 and len(chunk_words) > 1:
            # Find overlap position by counting characters backwards
            overlap_chars = 0
            overlap_words = 0
            
            for i in range(len(chunk_words) - 1, -1, -1):
                word_length = len(chunk_words[i]) + (1 if i < len(chunk_words) - 1 else 0)
                if overlap_chars + word_length > overlap:
                    break
                overlap_chars += word_length
                overlap_words += 1
            
            # Move start index forward, accounting for overlap
            next_start = start_idx + len(chunk_words) - overlap_words
            
            # Ensure we make progress
            if next_start <= start_idx:
                next_start = start_idx + max(1, len(chunk_words) // 2)
        else:
            next_start = start_idx + len(chunk_words)
        
        start_idx = next_start
        
        # Break if we've processed all words
        if start_idx >= len(words):
            break
    
    logger.debug(f"Chunked text into {len(chunks)} chunks (size: {chunk_size}, overlap: {overlap})")
    return chunks

def get_file_parser(file_path: Path) -> Optional[callable]:
    """
    Get appropriate parser function for a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        callable: Parser function or None if unsupported format
    """
    suffix = file_path.suffix.lower()
    
    parsers = {
        '.txt': parse_txt,
        '.md': parse_txt,
        '.text': parse_txt,
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.doc': parse_docx,  # Note: .doc files might not work with python-docx
        '.png': parse_image_ocr,
        '.jpg': parse_image_ocr,
        '.jpeg': parse_image_ocr,
        '.gif': parse_image_ocr,
        '.bmp': parse_image_ocr,
        '.tiff': parse_image_ocr,
    }
    
    return parsers.get(suffix)

def parse_file(file_path: Path) -> str:
    """
    Parse any supported file and return its text content.
    Handles errors gracefully and returns empty string for problematic files.
    
    Args:
        file_path: Path to the file to parse
        
    Returns:
        str: Extracted text content, or empty string if parsing fails
    """
    try:
        parser = get_file_parser(file_path)
        if parser is None:
            print(f"⚠️  Unsupported file format: {file_path.suffix} for {file_path.name}")
            logger.warning(f"Unsupported file format: {file_path.suffix} for file {file_path}")
            return ""
        
        # Special handling for images and PDFs to show OCR status
        if parser == parse_image_ocr:
            print(f"🖼️  Detected image file: {file_path.name} - Using OCR")
        elif parser == parse_pdf:
            print(f"📄 Detected PDF file: {file_path.name} - Will use OCR for image-based pages")
            
        result = parser(file_path)
        
        # Ensure we return a string
        if result is None:
            return ""
        
        parsed_result = str(result).strip()
        
        # Show success status for OCR specifically
        if parser == parse_image_ocr:
            if parsed_result:
                print(f"✅ OCR completed successfully for {file_path.name}")
            else:
                print(f"⚠️  OCR completed but no text extracted from {file_path.name}")
        elif parser == parse_pdf:
            if parsed_result:
                print(f"✅ PDF processing completed for {file_path.name} ({len(parsed_result)} characters)")
            else:
                print(f"⚠️  PDF processing completed but no text extracted from {file_path.name}")
        
        return parsed_result
        
    except Exception as e:
        print(f"❌ Error parsing file {file_path.name}: {str(e)}")
        logger.error(f"Error parsing file {file_path}: {e}")
        # Return empty string instead of raising exception
        return ""
