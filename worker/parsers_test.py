#!/usr/bin/env python3
"""
Test suite for parsers.py functionality.

This script tests document parsing and text chunking capabilities.
"""

import tempfile
import logging
from pathlib import Path
from typing import List

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import parsers

def test_chunk_text():
    """Test the text chunking functionality."""
    print("=== Testing Text Chunking ===\n")
    
    # Test case 1: Basic chunking
    print("1. Basic chunking test:")
    text = "This is a test document with several words that should be chunked properly."
    chunks = parsers.chunk_text(text, chunk_size=30, overlap=10)
    
    print(f"   Original text: {text}")
    print(f"   Chunk size: 30, Overlap: 10")
    print(f"   Number of chunks: {len(chunks)}")
    for i, chunk in enumerate(chunks, 1):
        print(f"   Chunk {i}: '{chunk}' ({len(chunk)} chars)")
    
    # Test case 2: Longer text
    print("\n2. Longer text chunking:")
    long_text = " ".join([
        "The quick brown fox jumps over the lazy dog.",
        "This is a longer sentence that contains more words and should be split across multiple chunks.",
        "Python is a powerful programming language that is widely used for data science and machine learning.",
        "Vector embeddings are numerical representations of text that capture semantic meaning."
    ])
    
    chunks = parsers.chunk_text(long_text, chunk_size=100, overlap=20)
    print(f"   Original length: {len(long_text)} chars")
    print(f"   Chunk size: 100, Overlap: 20")
    print(f"   Number of chunks: {len(chunks)}")
    
    for i, chunk in enumerate(chunks, 1):
        print(f"   Chunk {i}: '{chunk[:50]}...' ({len(chunk)} chars)")
    
    # Test case 3: Edge cases
    print("\n3. Edge cases:")
    
    # Empty text
    empty_chunks = parsers.chunk_text("", 50, 10)
    print(f"   Empty text chunks: {len(empty_chunks)}")
    
    # Single word
    single_chunks = parsers.chunk_text("Hello", 50, 10)
    print(f"   Single word chunks: {len(single_chunks)} - {single_chunks}")
    
    # Text shorter than chunk size
    short_chunks = parsers.chunk_text("Short text", 100, 20)
    print(f"   Short text chunks: {len(short_chunks)} - {short_chunks}")
    
    # Test overlap validation
    try:
        parsers.chunk_text("test", 10, 15)  # overlap > chunk_size
        print("   ✗ Should have failed with overlap > chunk_size")
    except ValueError as e:
        print(f"   ✓ Correctly caught error: {e}")

def test_text_parsing():
    """Test text file parsing."""
    print("\n=== Testing Text File Parsing ===\n")
    
    # Create temporary text file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        test_content = """This is a test document.
It contains multiple lines.
And should be parsed correctly.

Even with empty lines in between."""
        f.write(test_content)
        temp_path = Path(f.name)
    
    try:
        parsed_content = parsers.parse_txt(temp_path)
        print(f"✓ Parsed text file successfully")
        print(f"  Original length: {len(test_content)}")
        print(f"  Parsed length: {len(parsed_content)}")
        print(f"  Content matches: {test_content.strip() == parsed_content}")
        print(f"  First 100 chars: {parsed_content[:100]}")
        
    except Exception as e:
        print(f"✗ Text parsing failed: {e}")
    finally:
        temp_path.unlink()  # Clean up

def test_pdf_parsing():
    """Test PDF parsing (if PyMuPDF is available)."""
    print("\n=== Testing PDF Parsing ===\n")
    
    try:
        # Create a simple PDF for testing (this is a basic test)
        import fitz
        
        # Create temporary PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = Path(f.name)
        
        # Create a simple PDF document
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is a test PDF document.\nIt has multiple lines.\nFor testing purposes.")
        doc.save(temp_path)
        doc.close()
        
        try:
            parsed_content = parsers.parse_pdf(temp_path)
            print(f"✓ Parsed PDF successfully")
            print(f"  Content length: {len(parsed_content)}")
            print(f"  Content: {parsed_content}")
            
        except Exception as e:
            print(f"✗ PDF parsing failed: {e}")
        finally:
            temp_path.unlink()  # Clean up
            
    except ImportError:
        print("⚠️  PyMuPDF not available, skipping PDF test")
    except Exception as e:
        print(f"✗ PDF test setup failed: {e}")

def test_docx_parsing():
    """Test DOCX parsing (if python-docx is available)."""
    print("\n=== Testing DOCX Parsing ===\n")
    
    try:
        from docx import Document
        
        # Create temporary DOCX
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = Path(f.name)
        
        # Create a simple DOCX document
        doc = Document()
        doc.add_paragraph("This is a test DOCX document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.add_paragraph("For comprehensive testing.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        
        doc.save(temp_path)
        
        try:
            parsed_content = parsers.parse_docx(temp_path)
            print(f"✓ Parsed DOCX successfully")
            print(f"  Content length: {len(parsed_content)}")
            print(f"  Content: {parsed_content}")
            
        except Exception as e:
            print(f"✗ DOCX parsing failed: {e}")
        finally:
            temp_path.unlink()  # Clean up
            
    except ImportError:
        print("⚠️  python-docx not available, skipping DOCX test")
    except Exception as e:
        print(f"✗ DOCX test setup failed: {e}")

def test_file_parser_detection():
    """Test automatic file format detection."""
    print("\n=== Testing File Format Detection ===\n")
    
    test_files = [
        "document.txt",
        "document.pdf", 
        "document.docx",
        "image.png",
        "readme.md",
        "unsupported.xyz"
    ]
    
    for filename in test_files:
        path = Path(filename)
        parser = parsers.get_file_parser(path)
        
        if parser:
            print(f"✓ {filename}: {parser.__name__}")
        else:
            print(f"✗ {filename}: No parser found")

def test_chunking_determinism():
    """Test that chunking is deterministic."""
    print("\n=== Testing Chunking Determinism ===\n")
    
    text = "This is a test document that should be chunked consistently every time we run the chunking algorithm on it."
    
    # Run chunking multiple times
    chunks1 = parsers.chunk_text(text, chunk_size=50, overlap=15)
    chunks2 = parsers.chunk_text(text, chunk_size=50, overlap=15)
    chunks3 = parsers.chunk_text(text, chunk_size=50, overlap=15)
    
    all_same = chunks1 == chunks2 == chunks3
    print(f"✓ Chunking is deterministic: {all_same}")
    
    if all_same:
        print(f"  Consistent result: {len(chunks1)} chunks")
        for i, chunk in enumerate(chunks1, 1):
            print(f"    Chunk {i}: '{chunk}' ({len(chunk)} chars)")
    else:
        print("✗ Chunking results differ between runs!")

def main():
    """Run all tests."""
    print("🧪 Testing parsers.py functionality\n")
    
    try:
        test_chunk_text()
        test_chunking_determinism()
        test_file_parser_detection()
        test_text_parsing()
        test_pdf_parsing()
        test_docx_parsing()
        
        print("\n✅ All tests completed!")
        
    except Exception as e:
        print(f"\n❌ Test suite failed: {e}")
        logger.exception("Test suite failed")

if __name__ == "__main__":
    main()
